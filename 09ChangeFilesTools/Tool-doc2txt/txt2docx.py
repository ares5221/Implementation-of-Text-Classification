#!/usr/bin/env python
# _*_ coding:utf-8 _*_
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import re
'''将一个目录下所有txx文件转成docx,文件名用txt中的name'''
txtpath = os.path.abspath('./txtfolderFordocx')
docxpath = os.path.abspath('./docxfolder/')

all_FileNum = 0


def Translate(path):
    global all_FileNum
    files = os.listdir(path)  # 该目录下所有文件的名字
    for f in files:
        if (f[0] == '~' or f[0] == '.'):
            continue
        filepath = path + '\\' + f
        if filepath[-4:] == '.txt':
            filename = f[:-4]
            print(filename)
            pattern_rule2 = re.compile(r'\n')
            with open(filepath, 'r', encoding='utf-8') as ff:  # 读取txt中内容
                txtcontent = ff.read()
                print('title-----------------', ff.readline())
                txtcontent = re.sub(pattern_rule2, '', txtcontent)
                print(txtcontent)

            document = Document()  # 新建docx文件
            p = document.add_paragraph(txtcontent)
            document.styles['Normal'].font.name = u'宋体'
            # p.add_run(txtcontent).font.size = Pt(10.5)
            document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            # p.add_run.font.size = Pt(36)
            newdocxpath = docxpath + filename + '.docx'
            print(newdocxpath)
            document.save(newdocxpath)
            # ph_format = paragraph.paragraph_format
            # ph_format.space_before = Pt(10)  # 设置段前间距
            # ph_format.space_after = Pt(12)  # 设置段后间距
            # ph_format.line_spacing = Pt(19)  # 设置行间距
            all_FileNum += 1


if __name__ == '__main__':
    Translate(txtpath)
    # Translate(path2)
    print('文件夹中文件转换完毕，文件总数 = ', all_FileNum)
